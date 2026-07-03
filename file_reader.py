# ======================================================
# File Reader
# ======================================================
#
# Turns files the user attaches in the chat into something
# the local model can use:
#
#   - Text / code / PDF / Word  ->  extracted as text context
#   - Images (.png .jpg ...)     ->  base64, sent to a VISION
#                                    model via Ollama's "images"
#
# Everything degrades gracefully - a bad file returns a short
# note instead of crashing the app.
# ======================================================

import os
import io
import base64

# ------------------------------------------------------
# Limits
# ------------------------------------------------------
# Local models have a limited context window, so cap how
# much text we pull from any single file.
# ------------------------------------------------------

MAX_CHARS_PER_FILE = 20000

# "Read the bytes as text" extensions.
TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown",
    ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".hpp",
    ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt",
    ".sh", ".bat", ".ps1",
    ".html", ".htm", ".css", ".xml", ".yaml", ".yml",
    ".json", ".csv", ".tsv", ".log", ".ini", ".cfg", ".toml",
    ".v", ".sv", ".vhd", ".vhdl",          # HDL files
}

# Images we hand to a vision model.
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}

# File types the uploader will accept (no leading dot).
ALLOWED_UPLOAD_TYPES = sorted(
    {ext.lstrip(".") for ext in TEXT_EXTENSIONS}
    | {ext.lstrip(".") for ext in IMAGE_EXTENSIONS}
    | {"pdf", "docx"}
)


# ======================================================
# Classification
# ======================================================

def _ext(uploaded_file):
    name = getattr(uploaded_file, "name", "file")
    return os.path.splitext(name)[1].lower()


def is_image(uploaded_file):
    return _ext(uploaded_file) in IMAGE_EXTENSIONS


def partition_uploads(uploaded_files):
    """Split a list of uploads into (text_like_files, image_files)."""
    text_files = []
    image_files = []

    for uploaded_file in uploaded_files or []:
        if is_image(uploaded_file):
            image_files.append(uploaded_file)
        else:
            text_files.append(uploaded_file)

    return text_files, image_files


# ======================================================
# Images  ->  base64 (for vision models)
# ======================================================

def encode_images(image_files):
    """
    Return a list of base64 strings, one per image, suitable for
    Ollama's message "images" field.
    """
    encoded = []

    for uploaded_file in image_files or []:
        try:
            data = uploaded_file.getvalue()
            encoded.append(base64.b64encode(data).decode("utf-8"))
        except Exception:
            # Skip anything unreadable rather than break the turn.
            continue

    return encoded


# ======================================================
# Text extraction
# ======================================================

def extract_text(uploaded_file):
    """
    Turn a single (non-image) UploadedFile into text.
    On failure returns a short "[Could not read ...]" note.
    """
    name = getattr(uploaded_file, "name", "file")
    ext = _ext(uploaded_file)

    try:
        data = uploaded_file.getvalue()
    except Exception as error:
        return f"[Could not read {name}: {error}]"

    if ext == ".pdf":
        text = _read_pdf(data, name)
    elif ext == ".docx":
        text = _read_docx(data, name)
    elif ext in TEXT_EXTENSIONS:
        text = _read_plain_text(data)
    else:
        return (
            f"[{name} is a '{ext}' file, which isn't supported yet. "
            "Supported: text/code files, PDF, Word (.docx), and images.]"
        )

    return _truncate(text, name)


def build_context(text_files):
    """
    Combine one or more text-like files into a single context
    block, each wrapped with a clear header. Returns "" if empty.
    """
    if not text_files:
        return ""

    blocks = []

    for uploaded_file in text_files:
        name = getattr(uploaded_file, "name", "file")
        content = extract_text(uploaded_file)

        blocks.append(
            f"----- START OF FILE: {name} -----\n"
            f"{content}\n"
            f"----- END OF FILE: {name} -----"
        )

    return "\n\n".join(blocks)


# ======================================================
# Format-specific readers
# ======================================================

def _read_plain_text(data):
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _read_pdf(data, name):
    try:
        from pypdf import PdfReader
    except ImportError:
        return (
            "[PDF support needs the 'pypdf' package. "
            "Install it with:  pip install pypdf]"
        )

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as error:
        return f"[Could not open PDF {name}: {error}]"

    pages = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            pages.append(f"[Page {number}]\n{page_text.strip()}")

    if not pages:
        return (
            f"[No selectable text found in {name}. "
            "It may be a scanned PDF (image only).]"
        )

    return "\n\n".join(pages)


def _read_docx(data, name):
    try:
        import docx
    except ImportError:
        return (
            "[Word support needs the 'python-docx' package. "
            "Install it with:  pip install python-docx]"
        )

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as error:
        return f"[Could not open Word file {name}: {error}]"

    parts = [
        p.text for p in document.paragraphs if p.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    if not parts:
        return f"[No text found in {name}.]"

    return "\n".join(parts)


# ======================================================
# Internal
# ======================================================

def _truncate(text, name):
    if len(text) <= MAX_CHARS_PER_FILE:
        return text
    keep = text[:MAX_CHARS_PER_FILE]
    return (
        f"{keep}\n\n"
        f"[...{name} was truncated here. "
        f"Showing the first {MAX_CHARS_PER_FILE} characters only.]"
    )
